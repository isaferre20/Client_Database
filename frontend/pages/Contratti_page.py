import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from backend.client_data_backend import app as flask_app
flask_app.app_context().push()
import base64
from datetime import datetime
import os
import pandas as pd
import streamlit as st
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode, JsCode
from st_aggrid.shared import DataReturnMode
from backend.client_data_backend import db, Contract, Client
import sys, os, shutil, re, unicodedata, subprocess, locale
from datetime import datetime
from pathlib import Path
import pandas as pd
import streamlit as st
from st_aggrid import AgGrid, GridOptionsBuilder, GridUpdateMode
from docx import Document
from backend.client_data_backend import app as flask_app, get_clients_by_search, db, Client, Contract

from pathlib import Path
import shutil

import pythoncom
from win32com.client import Dispatch

locale.setlocale(locale.LC_TIME, 'it_IT.UTF-8')
flask_app.app_context().push()

from frontend.navbar import navbar

navbar()

# Button to open IVA folder (top-right)
with st.container():
    col1, col2 = st.columns([2, 8])
    with col1:
        if st.button("📂 Apri Cartella CONTRATTI", key="open_iva_folder"):
            try:
                import subprocess
                iva_folder_path = "Z:\\Documents\\Lavori Idraulica\\Isa uso ufficio\\Client_DB\\CONTRATTI"
                if os.name == 'nt':
                    os.startfile(iva_folder_path)
                elif os.name == 'posix':
                    subprocess.Popen(["open", iva_folder_path])
            except Exception as e:
                st.error(f"❌ Errore nell'aprire la cartella CONTRATTI: {e}")

st.markdown(f"""
<style>
    html, body {{
        margin: 0;
        font-family: 'Segoe UI', sans-serif;
        background: radial-gradient(circle at top left, #fef6ff 0%, #eef3ff 30%, #f5faff 100%);
        background-attachment: fixed;
        background-repeat: no-repeat;
        background-size: cover;
        overflow-x: hidden;
        color: #1a3fc1;
    }}
    body::before {{
        content: "";
        position: fixed;
        top: -150px;
        left: -100px;
        width: 600px;
        height: 600px;
        background: radial-gradient(circle, rgba(200, 220, 255, 0.3), transparent 100%);
        z-index: -1;
        filter: blur(90px);
    }}
    body::after {{
        content: "";
        position: fixed;
        bottom: -140px;
        right: -120px;
        width: 600px;
        height: 600px;
        background: radial-gradient(circle, rgba(255, 230, 250, 0.3), transparent 100%);
        z-index: -1;
        filter: blur(90px);
    }}

    .container {{
        max-width: 1200px;
        margin: 0 auto;
        padding: 0 2rem;
    }}


    .stButton > button {{
        background: none !important;
        border: none !important;
        color: #1a3fc1 !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
        padding: 0.4rem 0.6rem !important;
    }}

    .stButton > button:hover {{
        color: #1430a5 !important;
        cursor: pointer;
    }}

    /* Universal override for Streamlit inputs */
    input[type="text"], input[type="email"], input[type="tel"], textarea, select {{
        background-color: white !important;
        color: #1a1a1a !important;
        border: 1px solid #d6d6d6 !important;
        border-radius: 10px !important;
        padding: 0.5rem 0.75rem !important;
        box-shadow: 0 1px 4px rgba(0, 0, 0, 0.05);
    }}
    [data-baseweb="input"] input,
    [data-baseweb="select"] {{
        background-color: white !important;
        color: #1a1a1a !important;
    }}
    input:focus, textarea:focus, select:focus {{
        outline: none !important;
        box-shadow: 0 0 0 2px rgba(26, 63, 193, 0.2) !important;
    }}

    h2 {{
        color: #1a3fc1;
        margin-bottom: 1.5rem;
    }}
    
    /* Selectbox selected value container background */
    div[data-baseweb="select"] > div > div {{
        background-color: white !important;
        color: #1a1a1a !important;
    }}

</style>
""", unsafe_allow_html=True)


TEMPLATE_PATHS = {
    "MANUTENZIONE (cm)": "Z:\\Documents\\Lavori Idraulica\\Isa uso ufficio\\Client_DB\\Client_Database\\templates_docs\\MANUT.docx",
    "Contratto 3RESP + MANUTENZIONE (crm)": "Z:\\Documents\\Lavori Idraulica\\Isa uso ufficio\\\\Client_DB\\Client_Database\\templates_docs\\MANUT_3RESP.docx"
}

# --- UTILS ---
def sanitize_filename(name):
    if not name: return "Senza_Nome"
    name = unicodedata.normalize('NFKD', name).encode('ASCII', 'ignore').decode('utf-8')
    name = re.sub(r'[^\w\s\-]', '', name)
    return name.strip().replace(' ', '_')

def validate_cf_piva(value): value = value.strip(); return (len(value)==16 and value.isalnum()) or (len(value)==11 and value.isdigit())
def is_valid_email(email): return bool(re.match(r".+@.+\..+", email.strip()))

def parse_date(value):
    if pd.isna(value):
        return None
    if isinstance(value, str):
        try:
            return datetime.strptime(value, "%d/%m/%Y")
        except ValueError:
            try:
                return datetime.fromisoformat(value)
            except:
                return None
    elif isinstance(value, datetime):
        return value
    return None


def format_date(date_value):
    try: return pd.to_datetime(date_value, dayfirst=True).strftime('%d %B %Y')
    except: return date_value or ""

def replace_placeholder(doc, placeholder, value):
    value = "" if not value or str(value) in ["nan", "None"] else str(value)
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, value)

def remove_block(doc, start, end, keep):
    in_block = False; to_del = []
    for para in doc.paragraphs:
        if start in para.text: in_block = True; (to_del.append(para) if not keep else para.text.replace(start,''))
        elif end in para.text: in_block = False; (to_del.append(para) if not keep else para.text.replace(end,''))
        elif in_block and not keep: to_del.append(para)
    for para in to_del: p = para._element; p.getparent().remove(p)

def conditional_block(doc, tot_pot):
    """
    Se tot_pot > 116: rimuove SOLO il tag {{IF_TOT_POT>116}} e {{END_IF}}, ma lascia il testo in mezzo.
    Se tot_pot <= 116: elimina TUTTO il blocco (tag + testo tra i tag).
    """
    start_tag = "{{IF_TOT_POT>116}}"
    end_tag = "{{END_IF}}"
    in_block = False
    to_delete = []
    for para in doc.paragraphs:
        if start_tag in para.text:
            if tot_pot > 116:
                para.text = para.text.replace(start_tag, "")
                in_block = True
            else:
                in_block = True
                to_delete.append(para)
        elif end_tag in para.text:
            if tot_pot > 116:
                para.text = para.text.replace(end_tag, "")
                in_block = False
            else:
                to_delete.append(para)
                in_block = False
        elif in_block and tot_pot <= 116:
            to_delete.append(para)
    # Rimuovi i paragrafi segnati
    for para in to_delete:
        p = para._element
        p.getparent().remove(p)


def get_next_contract_filename(prefix, year, ente):
    tag = f"{prefix}{year}"  # e.g., cm25
    all_contracts = Contract.query.all()

    pattern = re.compile(rf"{tag}(\d{{3}})")  # cm25 + 3 digits
    nums = []
    for c in all_contracts:
        if c.contract_url:
            match = pattern.search(Path(c.contract_url).stem)
            if match:
                nums.append(int(match.group(1)))

    next_number = max(nums, default=0) + 1
    return f"{tag}{next_number:03}_{sanitize_filename(ente)}"


def convert_to_pdf(docx_path, outdir):
    libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice"
    subprocess.run([libreoffice_path, "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(outdir)], check=True)

from pathlib import Path
import shutil

def mark_old_contract_superato(old_contract, client):
    old_path = Path(old_contract.contract_url)
    old_stem = old_path.stem
    if "_SUPERATO" not in old_stem:
        superato_stem = old_stem + "_SUPERATO"
        superato_docx = old_path.with_stem(superato_stem).with_suffix('.docx')
        superato_pdf = old_path.with_stem(superato_stem).with_suffix('.pdf')

        cliente_folder_name = f"{sanitize_filename(client.cognome)}_{sanitize_filename(client.nome)}_{client.codice_fiscale}"
        cliente_folder = Path("DOCUMENTAZIONE_CLIENTI") / cliente_folder_name / "CONTRATTI"
        cliente_folder.mkdir(parents=True, exist_ok=True)

        # PRIMA aggiorno i file linkati nella cartella cliente
        for ext in ['.pdf', '.docx']:
            old_file = old_path.with_suffix(ext)
            superato_file = old_path.with_stem(superato_stem).with_suffix(ext)
            link_path = cliente_folder / old_file.name
            new_link_path = cliente_folder / superato_file.name

            if link_path.exists():
                try:
                    link_path.rename(new_link_path)
                except Exception as e:
                    link_path.unlink()
                    if superato_file.exists():
                        shutil.copy2(superato_file, new_link_path)

        # POI aggiorno i file principali nella cartella CONTRATTI
        if old_path.exists():
            old_docx_path = old_path.with_suffix('.docx')
            if old_docx_path.exists():
                old_docx_path.rename(superato_docx)
            old_pdf_path = old_path.with_suffix('.pdf')
            if old_pdf_path.exists():
                old_pdf_path.rename(superato_pdf)
    old_contract.contract_url = str(superato_pdf)  # oppure str(superato_docx) se vuoi puntare al DOCX
    db.session.commit()

def generate_new_contract_filename(prefix, year, ente, old_suffix=None):
    base = get_next_contract_filename(prefix, year, ente)
    if old_suffix:
        # old_suffix tipo 'rev21001' o solo numerico? adattare di conseguenza
        return f"{base}_rev{old_suffix}"
    else:
        return base
    
from docx import Document

def replace_immobili_placeholder(doc, immobili):
    for i, para in enumerate(doc.paragraphs):
        if '{{IMMOBILI}}' in para.text:
            # Save the location
            p_element = para._element
            parent = p_element.getparent()
            index = parent.index(p_element)

            # Remove placeholder paragraph
            parent.remove(p_element)

            for imm_idx, immobile in enumerate(immobili):
                # Add title paragraph
                title = f"Immobile {imm_idx+1}: {immobile['indirizzo']} - Proprietà: {immobile['proprieta']} - Occupante: {immobile['occupante']}"
                new_para = doc.add_paragraph(title)
                parent.insert(index, new_para._element)
                index += 1

                # Add table
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Unità"
                hdr_cells[1].text = "Modello Installato"
                hdr_cells[2].text = "Potenza al Focolare"
                hdr_cells[3].text = "Data Installazione"
                hdr_cells[4].text = "DI.CO / DI.RI"
                hdr_cells[5].text = "Data DI.CO/DI.RI"
                hdr_cells[6].text = "Codice Impianto"

                for unita in immobile.get("unita", []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(unita.get("nome", ""))
                    row_cells[1].text = str(unita.get("modello", ""))
                    row_cells[2].text = str(unita.get("potenza_foc", ""))
                    row_cells[3].text = str(unita.get("data_installazione", ""))
                    row_cells[4].text = str(unita.get("dico_diri", ""))
                    row_cells[5].text = str(unita.get("data_dico_diri", ""))
                    row_cells[6].text = str(unita.get("codice_impianto", ""))

                # Insert table element
                parent.insert(index, table._element)
                index += 1
            break

# --- LAYOUT ---
st.title("📝 Gestione Contratti")
tabs = ["📄 Genera Contratto", "Visualizza Contratti"]
if "active_tab" not in st.session_state: st.session_state["active_tab"] = tabs[0]
selected_tab = st.radio("📂 Seleziona sezione", tabs, index=tabs.index(st.session_state["active_tab"]), horizontal=True)
st.session_state["active_tab"] = selected_tab

if selected_tab == tabs[0]:
    revise_id = st.session_state.get("revise_contract_id")
    preselect_client = None
    prefill_data = {}

    if revise_id:
        old_contract = db.session.get(Contract, revise_id)
        if old_contract:
            preselect_client = db.session.get(Client, old_contract.client_id)
            prefill_data = {
                "inizio_stagione": old_contract.inizio_stagione or "",
                "fine_stagione": old_contract.fine_stagione or "",
                "data_doc": datetime.now().strftime("%d/%m/%Y"),
                "contract_type": old_contract.contract_type
            }
    # --- Ricerca Cliente ---
    search_query = st.text_input("🔍 Cerca cliente (nome, cognome o codice fiscale)", value=preselect_client.codice_fiscale if preselect_client else "")
    if search_query:
        clients = get_clients_by_search(search_query)
        if clients:
            client = st.selectbox("👤 Seleziona Cliente", clients, format_func=lambda c: f"{c.nome} {c.cognome} ({c.codice_fiscale})")
            if client:
                model = st.selectbox("📄 Seleziona Modello di Contratto", ["", *TEMPLATE_PATHS], index=0)
                if model and model in TEMPLATE_PATHS:
                    # --- TABS LAYOUT ---
                    tab1, tab2, tab3, tab4 = st.tabs([
                        "📋 Dati Generali", 
                        "💶 Costi / Tariffe", 
                        "⚙️ Generatori" if model=="Contratto 3RESP + MANUTENZIONE (crm)" else "🏘️ Immobili", 
                        "📝 Riepilogo & Genera"
                    ])
                    with tab1:
                        ente = st.text_input("ENTE", value=f"{client.cognome} {client.nome}")
                        cf = st.text_input("P.IVA / Codice Fiscale", value=client.codice_fiscale)
                        email = st.text_input("Email", value=client.mail)
                        indirizzo = st.text_input("Indirizzo", value=f"{client.indirizzo_residenza}, {client.civico} - {client.citta_residenza} ({client.provincia})")
                        responsabile = st.text_input("Responsabile Impianto", value=f"{client.nome} {client.cognome}")
                        cf_resp = st.text_input("CF Responsabile Impianto", value=client.codice_fiscale)
                        legale = st.text_input("Legale Rappresentante", value=f"{client.nome} {client.cognome}")
                        cf_legale = st.text_input("CF Legale Rappresentante", value=client.codice_fiscale)
                        direttore = st.text_input("Direttore / Direttrice (Opzionale)")
                        luogo = st.text_input("Luogo Firma", value="Carmagnola")
                        data_doc = st.text_input("Data Firma", value=datetime.now().strftime('%d/%m/%Y'))
                        inizio_stagione = st.text_input("Inizio Stagione", value=f"15/10/{datetime.now().year}")
                        fine_stagione = st.text_input("Fine Stagione", value=f"15/04/{datetime.now().year + 1}")
                        st.session_state["input_generali"] = dict(
                            ente=ente, cf=cf, email=email, indirizzo=indirizzo,
                            responsabile=responsabile, cf_resp=cf_resp, legale=legale,
                            cf_legale=cf_legale, direttore=direttore, luogo=luogo,
                            data_doc=data_doc, inizio_stagione=inizio_stagione, fine_stagione=fine_stagione
                        )
                    with tab2:
                        euro_ord = st.number_input("Manutenzione Ordinaria (€)", value=750, step=50)
                        chiamata_feriali = st.number_input("Chiamate feriali (€)", value=35, step=5)
                        chiamata_festivi = st.number_input("Chiamate festivi (€)", value=60, step=5)
                        manodopera = st.number_input("Manodopera ordinaria (€ / ora)", value=35, step=5)
                        euro_resp, tipologia_stagione = 0, ""
                        if model == "Contratto 3RESP + MANUTENZIONE (crm)":
                            euro_resp = st.number_input("Terzo Responsabile (€)", value=800, step=50)
                            tipologia_stagione = st.selectbox("Tipologia Stagione", ["RISCALDAMENTO", "ACQUA SANITARIA", "RISCALDAMENTO + ACQUA SANITARIA"])
                        st.session_state["input_costi"] = dict(
                            euro_ord=euro_ord, chiamata_feriali=chiamata_feriali,
                            chiamata_festivi=chiamata_festivi, manodopera=manodopera,
                            euro_resp=euro_resp, tipologia_stagione=tipologia_stagione
                        )
                    with tab3:
                        generatori, total_pot = [], 0
                        if model == "Contratto 3RESP + MANUTENZIONE (crm)":
                            num_generators = st.number_input('Numero di generatori', min_value=1, step=1)
                            for i in range(int(num_generators)):
                                st.subheader(f'Generatore {i + 1}')
                                gen_data = {
                                    'Generatore': st.text_input(f'Numero Generatore {i+1}', value=str(i+1), key=f'gen_num_{i}'),
                                    'Marca e Modello': st.text_input(f'Marca e Modello {i+1}', key=f'gen_mod_{i}'),
                                    'Combustibile': st.text_input(f'Combustibile {i+1}', key=f'gen_comb_{i}'),
                                    'Potenza (kW)': st.number_input(f'Potenza (kW) {i+1}', min_value=0.0, step=0.1, key=f'gen_pot_{i}')
                                }
                                total_pot += gen_data['Potenza (kW)']
                                generatori.append(gen_data)
                        st.session_state["generatori"] = generatori
                        st.session_state["total_pot"] = total_pot
                        if model == "MANUTENZIONE (cm)":
                            # -- Inizializza immobili se non già presenti --
                            if "immobili" not in st.session_state:
                                st.session_state["immobili"] = []

                            if st.button("➕ Aggiungi Immobile"):
                                st.session_state["immobili"].append({
                                    "indirizzo": "",
                                    "proprieta": "",
                                    "occupante": "",
                                    "unita": []
                                })

                            for idx_imm, immobile in enumerate(st.session_state["immobili"]):
                                st.subheader(f"Immobile #{idx_imm + 1}")
                                immobile["indirizzo"] = st.text_input(f"Indirizzo immobile #{idx_imm+1}", value=immobile["indirizzo"], key=f"indirizzo_{idx_imm}")
                                immobile["proprieta"] = st.text_input(f"Proprietà immobile #{idx_imm+1}", value=immobile["proprieta"], key=f"proprieta_{idx_imm}")
                                immobile["occupante"] = st.text_input(f"Occupante immobile #{idx_imm+1}", value=immobile["occupante"], key=f"occupante_{idx_imm}")

                                # Bottone per aggiungere unità a questo immobile
                                if st.button(f"➕ Aggiungi Unità a immobile #{idx_imm + 1}", key=f"add_unita_{idx_imm}"):
                                    immobile["unita"].append({
                                        "nome": "",
                                        "modello": "",
                                        "potenza_foc": "",
                                        "data_installazione": "",
                                        "dico_diri": "",
                                        "data_dico_diri": "",
                                        "codice_impianto": ""
                                    })

                                for idx_unita, unita in enumerate(immobile["unita"]):
                                    with st.expander(f"Unità #{idx_unita+1} per Immobile #{idx_imm+1}", expanded=True):
                                        unita["nome"] = st.text_input(f"Nome unità #{idx_unita+1}", value=unita["nome"], key=f"nome_unita_{idx_imm}_{idx_unita}")
                                        unita["modello"] = st.text_input(f"Modello installato #{idx_unita+1}", value=unita["modello"], key=f"modello_unita_{idx_imm}_{idx_unita}")
                                        unita["potenza_foc"] = st.text_input(f"Potenza al focolare #{idx_unita+1}", value=unita["potenza_foc"], key=f"potenza_foc_unita_{idx_imm}_{idx_unita}")
                                        unita["data_installazione"] = st.text_input(f"Data installazione #{idx_unita+1}", value=unita["data_installazione"], key=f"data_installazione_{idx_imm}_{idx_unita}")
                                        unita["dico_diri"] = st.text_input(f"DI.CO / DI.RI #{idx_unita+1}", value=unita["dico_diri"], key=f"dico_diri_unita_{idx_imm}_{idx_unita}")
                                        unita["data_dico_diri"] = st.text_input(f"Data DI.CO / DI.RI #{idx_unita+1}", value=unita["data_dico_diri"], key=f"data_dico_diri_unita_{idx_imm}_{idx_unita}")
                                        unita["codice_impianto"] = st.text_input(f"Codice impianto #{idx_unita+1}", value=unita["codice_impianto"], key=f"codice_impianto_unita_{idx_imm}_{idx_unita}")

                                        if st.button(f"🗑️ Rimuovi Unità #{idx_unita+1}", key=f"remove_unita_{idx_imm}_{idx_unita}"):
                                            immobile["unita"].pop(idx_unita)
                                            st.experimental_rerun()

                                if st.button(f"🗑️ Rimuovi Immobile #{idx_imm+1}", key=f"remove_immobile_{idx_imm}"):
                                    st.session_state["immobili"].pop(idx_imm)
                                    st.experimental_rerun()

                            st.markdown("---")
                            st.write("**Riepilogo immobili/unità compilati:**")
                            st.json(st.session_state["immobili"])
                            # Al salvataggio, passa questa lista dove vuoi (DB, documento Word, ecc)
                            st.session_state["immobili_salvati"] = st.session_state["immobili"]
                    with tab4:
                        # --- RECUPERA VALORI ---
                        dati = {**st.session_state.get("input_generali", {}), **st.session_state.get("input_costi", {})}
                        if model == "Contratto 3RESP + MANUTENZIONE (crm)":
                            dati["generatori"] = st.session_state.get("generatori", [])
                            dati["total_pot"] = st.session_state.get("total_pot", 0)
                        st.write("**Riepilogo dati:**")
                        st.json(dati)
                        if st.button("📄 Genera Contratto!"):
                            # --- VALIDAZIONE ---
                            if not all([dati.get('ente'), dati.get('cf'), dati.get('email'), dati.get('indirizzo')]): st.error("❗ Compila tutti i campi obbligatori."); st.stop()
                            if not validate_cf_piva(dati.get('cf')): st.error("CF/PIVA non valido"); st.stop()
                            if not is_valid_email(dati.get('email')): st.error("Email non valida"); st.stop()
                            # --- CREAZIONE ---
                            doc = Document(TEMPLATE_PATHS[model])
                            placeholders = {
                                'ENTE': dati.get('ente'), 'P.IVA_CF': dati.get('cf'), 'EMAIL': dati.get('email'),
                                'INDIRIZZO': dati.get('indirizzo'), 'Responsabile_Impianto': dati.get('responsabile'),
                                'CF_Resp_Impianto': dati.get('cf_resp'), 'Legale_rappresentante': dati.get('legale'),
                                'CF_L_Rappresentante': dati.get('cf_legale'), 'Direttore_Direttrice': dati.get('direttore'),
                                'INIZIO_STAGIONE': dati.get('inizio_stagione'), 'FINE_STAGIONE': dati.get('fine_stagione'),
                                'LUOGO': dati.get('luogo'), 'DATA': dati.get('data_doc'),
                            }
                            if model == "MANUTENZIONE (cm)":
                                placeholders.update({
                                    '€€_ORD': dati.get('euro_ord'), '€€_FERIALI': dati.get('chiamata_feriali'),
                                    '€€_FESTIVI': dati.get('chiamata_festivi'), '€€_MO': dati.get('manodopera')
                                })
                            if model == "Contratto 3RESP + MANUTENZIONE (crm)":
                                placeholders.update({
                                    'TIPOLOGIA_STAGIONE': dati.get('tipologia_stagione'), 'EURO_MANUTENZIONE': dati.get('euro_ord'),
                                    'EURO_TERZO_RESP': dati.get('euro_resp'), 'TOTALE': dati.get('euro_ord',0)+dati.get('euro_resp',0),
                                    'CHIAMATA_FERIALI': dati.get('chiamata_feriali'), 'CHIAMATA_FESTIVI': dati.get('chiamata_festivi'),
                                    'MANODOPERA': dati.get('manodopera'), 'TOT_POT': dati.get('total_pot',0),
                                    'SCADENZA': (pd.to_datetime(dati.get('data_doc'), dayfirst=True) + pd.DateOffset(years=1)).strftime('%d/%m/%Y')
                                })
                            for k,v in placeholders.items(): replace_placeholder(doc, f"{{{{{k}}}}}", v)
                            if model == "Contratto 3RESP + MANUTENZIONE (crm)": conditional_block(doc, dati.get('total_pot',0))
                            if model == "Contratto 3RESP + MANUTENZIONE (crm)" and dati.get('generatori'):
                                for table in doc.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            if '{{TAB_DATA}}' in cell.text:
                                                cell.text = ''
                                                nested_table = cell.add_table(rows=1, cols=len(dati['generatori'][0]))
                                                nested_table.style = 'Table Grid'
                                                for idx, col_name in enumerate(dati['generatori'][0].keys()):
                                                    nested_table.cell(0, idx).text = str(col_name)
                                                for row_data in dati['generatori']:
                                                    row_cells = nested_table.add_row().cells
                                                    for idx, value in enumerate(row_data.values()):
                                                        row_cells[idx].text = str(value)
                            if model == "MANUTENZIONE (cm)":
                                replace_immobili_placeholder(doc, st.session_state.get("immobili", []))
                            # --- Salvataggio file ---
                            year = pd.to_datetime(dati.get('data_doc'), dayfirst=True).strftime('%y')
                            folder = Path("Z:/Documents/Lavori Idraulica/Isa uso ufficio/Client_DB/CONTRATTI") / f"{sanitize_filename(client.cognome)}_{sanitize_filename(client.nome)}_{client.codice_fiscale}"
                            folder.mkdir(parents=True, exist_ok=True)
                            prefix = "cm" if model=="MANUTENZIONE (cm)" else "crm"
                            filename_code = get_next_contract_filename(prefix, year, f"{client.cognome}_{client.nome}_{client.codice_fiscale}")
                            docx_path = folder / f"{filename_code}.docx"
                            pdf_path = docx_path.with_suffix(".pdf")

                            if revise_id:
                                old_contract = db.session.get(Contract, revise_id)
                                if old_contract:
                                    mark_old_contract_superato(old_contract, client)

                                    old_path = Path(old_contract.contract_url)
                                    old_code = old_path.stem
                                    old_suffix = None
                                    match = re.search(r'rev(\d+)', old_code)
                                    if match:
                                        old_suffix = match.group(1)
                                    else:
                                        digits = re.findall(r'\d+', old_code)
                                        if digits:
                                            old_suffix = digits[-1]

                                    filename_code = f"{get_next_contract_filename(prefix, year, f"{client.cognome}_{client.nome}_{client.codice_fiscale}")}_rev{old_suffix}"
                                else:
                                    filename_code = get_next_contract_filename(prefix, year, f"{client.cognome}_{client.nome}_{client.codice_fiscale}")
                            else:
                                filename_code = get_next_contract_filename(prefix, year, f"{client.cognome}_{client.nome}_{client.codice_fiscale}")
                            

                            docx_path = folder / f"{filename_code}.docx"
                            pdf_path = docx_path.with_suffix(".pdf")

                            doc.save(docx_path)
                            convert_to_pdf(docx_path, folder)

                            # Usa lo stesso path corretto per DOCUMENTAZIONE_CLIENTI
                            cliente_folder_name = f"{sanitize_filename(client.cognome)}_{sanitize_filename(client.nome)}_{client.codice_fiscale}"
                            link_folder = Path("Z:/Documents/Lavori Idraulica/Isa uso ufficio/Client_DB/DOCUMENTAZIONE_CLIENTI") / cliente_folder_name / "CONTRATTI"
                            link_folder.mkdir(parents=True, exist_ok=True)
                            link_path = link_folder / pdf_path.name

                            if link_path.exists() or link_path.is_symlink():
                                link_path.unlink()

                            import pythoncom
                            from win32com.client import Dispatch

                            def create_windows_shortcut(target, shortcut_path):
                                pythoncom.CoInitialize()  # ✅ Required to use COM
                                shell = Dispatch('WScript.Shell')
                                shortcut = shell.CreateShortcut(str(shortcut_path))
                                shortcut.TargetPath = str(target)
                                shortcut.WorkingDirectory = str(target.parent)
                                shortcut.save()
                            # inside your logic
                            shortcut_path = link_folder / (pdf_path.stem + ".lnk")
                            create_windows_shortcut(pdf_path.resolve(), shortcut_path)


                            # salva in DB...
                            new_contract = Contract(
                                contract_url=str(pdf_path), contract_type=model,
                                contract_date=pd.to_datetime(dati.get('data_doc'), dayfirst=True).date(),
                                client_id=client.id, inizio_stagione=dati.get('inizio_stagione'),
                                fine_stagione=dati.get('fine_stagione'),
                                tot_ft=(dati.get('euro_ord',0)+(dati.get('euro_resp',0) if model=="Contratto 3RESP + MANUTENZIONE (crm)" else 0)),
                                data_prox_ft=(pd.to_datetime(dati.get('data_doc'), dayfirst=True) + pd.DateOffset(years=1)).strftime('%d/%m/%Y')
                            )
                            db.session.add(new_contract)
                            db.session.commit()
                            db.session.close()
                            st.success(f"✅ Contratto generato e salvato: {pdf_path.name}")

        else: st.warning("⚠ Nessun cliente trovato.")

if selected_tab == tabs[1]:
    st.markdown("### 📄 Visualizza Contratti")

    # 1) Load & normalize contracts
    try:
        contracts = (
        db.session.query(Contract, Client)
                 .join(Client, Contract.client_id == Client.id)
                 .all()
        )
    finally:
        db.session.close()

    rows = []
    for c, client in contracts:
        # normalize data_ultima_ft into a Python date or None
        raw = c.data_ultima_ft
        ultima_dt = None
        if isinstance(raw, datetime):
            ultima_dt = raw.date()
        elif isinstance(raw, str) and raw:
            try:
                # try ISO first
                ultima_dt = datetime.fromisoformat(raw).date()
            except ValueError:
                try:
                    # fallback dd/mm/yyyy
                    ultima_dt = datetime.strptime(raw, "%d/%m/%Y").date()
                except ValueError:
                    ultima_dt = None
            finally:
                db.session.close()

        # compute next‐year as a Python date or None
        prossima_dt = ultima_dt.replace(year=ultima_dt.year + 1) if ultima_dt else None

        rows.append({
            "ID": c.id,
            "Cliente": f"{client.nome} {client.cognome}",
            "Documento": os.path.basename(c.contract_url) if c.contract_url else "",
            "Inizio Stagione": c.inizio_stagione or "",
            "Fine Stagione": c.fine_stagione or "",
            "Totale Fatture (€)": f"{c.tot_ft or 0.0:.2f} €",
            # feed real dates
            "Data Ultima FT": ultima_dt.isoformat() if ultima_dt else "",
            "Data Prossima FT": prossima_dt.isoformat() if prossima_dt else "",
            "Tipo Contratto": c.contract_type,
            "Data Contratto": c.contract_date.strftime("%d %B %Y") if c.contract_date else "",
            "Percorso File": c.contract_url or ""
        })

    # 2) Build DataFrame
    if rows:
        df = pd.DataFrame(rows)[[
            "ID", "Cliente", "Documento", "Inizio Stagione", "Fine Stagione",
            "Totale Fatture (€)", "Data Ultima FT", "Data Prossima FT",
            "Tipo Contratto", "Data Contratto", "Percorso File"
        ]]
    else:
        df = pd.DataFrame(columns=[
            "ID", "Cliente", "Documento", "Inizio Stagione", "Fine Stagione",
            "Totale Fatture (€)", "Data Ultima FT", "Data Prossima FT",
            "Tipo Contratto", "Data Contratto", "Percorso File"
        ])


    # 3) Configure Ag-Grid
    gb = GridOptionsBuilder.from_dataframe(df)
    gb.configure_pagination()
    gb.configure_default_column(resizable=True, filter=True)  
    gb.configure_column("ID", headerCheckboxSelection=True, checkboxSelection=True)
    gb.configure_column("Percorso File", hide=True)

    # 📄 clickable link for Documento
    gb.configure_column(
        "Documento",
        headerName="📄 Contratto",
        cellRenderer=JsCode("""
            function(params) {
                return params.value || "";
            }
        """)
    )

    # shared date‐column params
    # Enable checkbox selection on rows
    gb.configure_selection(selection_mode="multiple", use_checkbox=True)

    # Data Ultima FT - date picker
    # Data Ultima FT
    gb.configure_column(
        "Data Ultima FT",
        headerName="📅 Data Ultima FT",
        editable=True,
        cellEditor="agDateCellEditor",
        cellEditorPopup=True,
        valueGetter=JsCode("""
            function(params) {
                const v = params.data["Data Ultima FT"];
                return v ? new Date(v) : null;
            }
        """),
        valueFormatter=JsCode("""
            function(params) {
                if (!params.value) return '';
                const d = new Date(params.value);
                return `${String(d.getDate()).padStart(2,'0')}/${String(d.getMonth()+1).padStart(2,'0')}/${d.getFullYear()}`;
            }
        """),
        valueSetter=JsCode("""
            function(params) {
                let d = params.newValue;
                if (typeof d === 'string') d = new Date(Date.parse(d));
                if (!(d instanceof Date) || isNaN(d)) return false;

                const yyyy = d.getFullYear();
                const mm = String(d.getMonth() + 1).padStart(2, '0');
                const dd = String(d.getDate()).padStart(2, '0');
                params.data["Data Ultima FT"] = `${yyyy}-${mm}-${dd}`;

                const dpft_raw = params.data["Data Prossima FT"];
                const dpft = new Date(dpft_raw);
                if (!dpft_raw || !(dpft instanceof Date) || isNaN(dpft)) {
                    const next = new Date(d);
                    next.setFullYear(next.getFullYear() + 1);
                    const yyyy2 = next.getFullYear();
                    const mm2 = String(next.getMonth() + 1).padStart(2, '0');
                    const dd2 = String(next.getDate()).padStart(2, '0');
                    params.data["Data Prossima FT"] = `${yyyy2}-${mm2}-${dd2}`;
                }

                return true;
            }
        """),
        filter="agDateColumnFilter"
    )

    # Data Prossima FT
    gb.configure_column(
    "Data Prossima FT",
    headerName="📅 Data Prossima FT",
    editable=True,
    cellEditor="agDateCellEditor",
    cellEditorPopup=True,
    valueGetter=JsCode("""
        function(params) {
            const v = params.data["Data Prossima FT"];
            return v ? new Date(v) : null;
        }
    """),
    valueFormatter=JsCode("""
        function(params) {
            if (!params.value) return '';
            const d = new Date(params.value);
            return `${String(d.getDate()).padStart(2,'0')}/${String(d.getMonth()+1).padStart(2,'0')}/${d.getFullYear()}`;
        }
    """),
    valueSetter=JsCode("""
        function(params) {
            let d = params.newValue;
            if (typeof d === 'string') d = new Date(Date.parse(d));
            if (!(d instanceof Date) || isNaN(d)) return false;

            const yyyy = d.getFullYear();
            const mm = String(d.getMonth() + 1).padStart(2, '0');
            const dd = String(d.getDate()).padStart(2, '0');
            params.data["Data Prossima FT"] = `${yyyy}-${mm}-${dd}`;

            return true;
        }
    """),
    filter="agDateColumnFilter"
)
    
    grid_opts = gb.build()

    # 4) Render the grid
    grid_resp = AgGrid(
        df,
        gridOptions=grid_opts,
        update_mode=GridUpdateMode.MODEL_CHANGED,
        data_return_mode=DataReturnMode.FILTERED_AND_SORTED,
        theme="streamlit",
        allow_unsafe_jscode=True,
        fit_columns_on_grid_load=True,
        rowSelection="multiple",  # ✅ allow multi-row selection
        use_checkbox=True,        # ✅ show checkboxes in the first column
    )

    # 5) Pull back as DataFrame
    raw = grid_resp["data"]
    updated_df = pd.DataFrame(raw)

    # 6) Save button → parse & commit
    if st.button("💾 Salva Modifiche"):
        def parse_any_date(val):
            if pd.isna(val) or not val:
                return None
            if isinstance(val, datetime):
                return val
            if isinstance(val, str):
                # ✅ Aggiorna l'ordine: metti prima il formato corretto "YYYY-MM-DD"
                for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%Y-%m-%dT%H:%M:%S.%fZ"):
                    try:
                        return datetime.strptime(val, fmt)
                    except ValueError:
                        continue
            return None


        updated_df = pd.DataFrame(grid_resp["data"])
        updated_df["_ultima_ts"] = updated_df["Data Ultima FT"].apply(parse_any_date)
        updated_df["_prossima_ts"] = updated_df["Data Prossima FT"].apply(parse_any_date)

        errors = []
        updated_count = 0

        try:
            db.session.expire_all()  # Ensures session fetches fresh data

            for _, row in updated_df.iterrows():
                cid = int(row["ID"])
                cont = db.session.get(Contract, cid)
                if not cont:
                    errors.append(f"⚠️ Contract ID {cid} not found.")
                    continue

                ultima = row["_ultima_ts"]
                prossima = row["_prossima_ts"]

                cont.data_ultima_ft = ultima.date() if pd.notna(ultima) else None
                cont.data_prox_ft   = prossima.date() if pd.notna(prossima) else None
                updated_count += 1

            db.session.commit()
            st.success(f"✅ Salvati {updated_count} contratto/i correttamente.")
            st.rerun()
            

        except Exception as e:
            db.session.rollback()
            st.error(f"❌ Errore durante il salvataggio: {e}")

        for e in errors:
            st.warning(e)


    # --- 9) Per-row actions on selected row (unchanged) ---
    sel = grid_resp.get("selected_rows")
    if isinstance(sel, pd.DataFrame):
        sel = sel.to_dict("records")
    if sel:
        selected_contract = sel[0]
        file_path = selected_contract.get("Percorso File", "").strip()
        r = sel[0]
        st.info(f"📑 {r['Cliente']} – {r['Documento']}")
        c1, c2 = st.columns(2)
        try:
            with c1:
                if st.button("✅ Registra Fattura Oggi", key=f"ft_oggi_{r['ID']}"):
                        c = db.session.get(Contract, int(r["ID"]))
                        if c:
                            today = datetime.now().date()
                            c.data_ultima_ft = today
                            c.data_prox_ft   = today.replace(year=today.year + 1)
                            db.session.commit()
                            st.success("✅ Data Ultima FT aggiornata a oggi.")
                            st.rerun()
                    
            with c2:
                if st.button("🔁 Modifica / Sostituisci Contratto", key=f"modifica_{r['ID']}"):
                    st.session_state["revise_contract_id"] = int(r["ID"])
                    st.session_state["active_tab"] = tabs[0]
                    st.rerun()
            if file_path and file_path.lower().endswith(".pdf") and Path(file_path).is_file():
                with open(file_path, "rb") as f:
                    base64_pdf = base64.b64encode(f.read()).decode("utf-8")

                st.markdown(f"""
                    <iframe src="data:application/pdf;base64,{base64_pdf}" 
                            width="100%" height="800px" type="application/pdf"></iframe>
                """, unsafe_allow_html=True)
            else:
                st.warning(f"⚠️ File non trovato o non è un PDF: {file_path}")
        finally:
                    db.session.close()

from docxtpl import DocxTemplate
import os
from datetime import date
from pathlib import Path
from docx import Document
import pandas as pd



def get_client_folder_path(client, base_dir="DOCUMENTAZIONE_CLIENTI"):
    client_id = client.get("id", "NOID")
    folder_name = f"{client['surname'].upper()}_{client['name'].upper()}_{client['codice_fiscale']}"
    client_path = os.path.join(base_dir, folder_name)
    os.makedirs(client_path, exist_ok=True)
    return client_path

def generate_doc_iva10(client, tipo_intervento, data_intervento, intervento, template_path="templates_docs/DICH_IVA_10.docx"):
    from docxtpl import DocxTemplate
    import os
    from datetime import date

    doc = DocxTemplate(template_path)
    O = "X" if tipo_intervento == "ordinaria" else ""
    S = "X" if tipo_intervento == "straordinaria" else ""

    if data_intervento is None:
        data_intervento = date.today().strftime("%d/%m/%Y")

    context = {
        "NOME": client["name"],
        "COGNOME": client["surname"],
        "CODICE_FISCALE": client["codice_fiscale"],
        "INDIRIZZO": client["address"],
        "NUM": client["address_number"],
        "CITTA": client["city"],
        "PROV": client["province"],
        "DATA_INTERVENTO": data_intervento,
        "O": O,
        "S": S,
        "INDIRIZZO_INT": intervento["indirizzo"],
        "NUM_INT": intervento["numero"],
        "CITTA_INT": intervento["citta"],
        "PROV_INT": intervento["provincia"]
    }

    doc.render(context)

    return doc

def generate_doc_iva4(client, intervento, titolo_abitativo, data_titolo, data_intervento, pratica, checkboxes, template_path="templates_docs/DICH_IVA_4_PRIMACASA.docx"):
    doc = DocxTemplate(template_path)
    
    if data_intervento is None:
        data_intervento = date.today().strftime("%d/%m/%Y")

    context = {
        "NOME": client["name"],
        "COGNOME": client["surname"],
        "CODICE_FISCALE": client["codice_fiscale"],
        "INDIRIZZO": client["address"],
        "NUM": client["address_number"],
        "CITTA": client["city"],
        "PROV": client["province"],
        "INDIRIZZO_INT": intervento["indirizzo_intervento"],
        "NUM_INT": intervento["civico_intervento"],
        "TIT_ABITATIVO": titolo_abitativo,
        "DATA_TIT": data_titolo,
        "TIT_INTERVENTO": intervento["citta_intervento"],
        "NUM_PRATICA": pratica["numero"],
        "COD_PRATICA": pratica["codice"],
        "DATA_INTERVENTO": data_intervento,
        "A": "X" if checkboxes.get("A") else "",
        "B": "X" if checkboxes.get("B") else "",
        "C": "X" if checkboxes.get("C") else "",
        "D": "X" if checkboxes.get("D") else "",
        "E": "X" if checkboxes.get("E") else "",
        "F": "X" if checkboxes.get("F") else "",
        "G": "X" if checkboxes.get("G") else ""
    }
    print("CONTEXT TO DOC:", context)

    doc.render(context)

    return doc

from docx import Document

from docx import Document

def generate_doc_dico(template_path: str, numero: str, data_doc: str, descrizione: str, legge: str, tipologia: str,
                      client_data: dict, intervento_data: dict):
    """
    Generates a DICO document by replacing placeholders in the template.

    Args:
        template_path (str): Path to the Word template.
        numero (str): DICO document number (e.g., "25001").
        data_doc (str): Document date (formatted dd/mm/yyyy).
        descrizione (str): Description of the intervention.
        legge (str): Legal references.
        tipologia (str): Type of intervention.
        client_data (dict): Data about the client.
        intervento_data (dict): Data about the intervention location.

    Returns:
        Document: The modified Word document object.
    """
    doc = Document(template_path)

    placeholders = {
        "{{NUM_DICO}}": numero,
        "{{DATA_DOC}}": data_doc,
        "{{DESCRIZIONE}}": descrizione,
        "{{LEGGE}}": legge,
        "{{TIPOLOGIA}}": tipologia,
        "{{NOME}}": client_data.get("nome", ""),
        "{{COGNOME}}": client_data.get("cognome", ""),
        "{{CODICE_FISCALE}}": client_data.get("codice_fiscale", ""),
        "{{INDIRIZZO}}": client_data.get("indirizzo", ""),
        "{{NUM}}": client_data.get("num", ""),
        "{{CITTA}}": client_data.get("citta", ""),
        "{{PROV}}": client_data.get("prov", ""),
        "{{INDIRIZZO_INT}}": intervento_data.get("indirizzo", ""),
        "{{NUM_INT}}": intervento_data.get("num", ""),
        "{{CITTA_INT}}": intervento_data.get("citta", ""),
        "{{PROV_INT}}": intervento_data.get("prov", ""),
        "{{FOGLIO}}": intervento_data.get("foglio", ""),
        "{{PART}}": intervento_data.get("part", ""),
        "{{SUB}}": intervento_data.get("sub", ""),
        "{{USO}}": intervento_data.get("uso", ""),
        "{{PROPR_NOME}}": client_data.get("nome", ""),
        "{{PROPR_COGNOME}}": client_data.get("cognome", "")
    }

    for paragraph in doc.paragraphs:
        for key, val in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    return doc

from docx import Document

def generate_contract_manut_3resp(data):
    doc = Document("templates_docs/MANUT_3RESP.docx")

    def replace_placeholders_in_doc(doc, replacements):
        def replace_in_paragraph(paragraph, replacements):
            if not hasattr(paragraph, "runs"):
                return
            full_text = ''.join(run.text for run in paragraph.runs)
            for key, val in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                full_text = full_text.replace(placeholder, str(val))
            # Remove all existing runs
            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            paragraph.add_run(full_text)

        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

    def remove_block(doc, start_marker, end_marker, keep_content=True):
        inside = False
        to_remove = []
        for para in doc.paragraphs:
            if start_marker in para.text:
                inside = True
                if keep_content:
                    para.text = para.text.replace(start_marker, "")
                else:
                    to_remove.append(para)
            elif end_marker in para.text:
                inside = False
                if keep_content:
                    para.text = para.text.replace(end_marker, "")
                else:
                    to_remove.append(para)
            elif inside and not keep_content:
                to_remove.append(para)

        for para in to_remove:
            p = para._element
            p.getparent().remove(p)

    def conditional_blocks(doc, potenza):
        try:
            pot = float(potenza)
        except (TypeError, ValueError):
            pot = 0
        if pot < 116:
            remove_block(doc, "{{IF_TOT_POT<116}}", "{{END_IF}}", keep_content=True)
            remove_block(doc, "{{IF_TOT_POT>116}}", "{{END_IF}}", keep_content=False)
        else:
            remove_block(doc, "{{IF_TOT_POT<116}}", "{{END_IF}}", keep_content=False)
            remove_block(doc, "{{IF_TOT_POT>116}}", "{{END_IF}}", keep_content=True)

    def insert_generator_table(doc, generators, placeholder):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = ""
                        nested = cell.add_table(rows=1, cols=4)
                        nested.style = "Table Grid"
                        headers = ["Generatore", "Marca e Modello", "Combustibile", "Potenza (kW)"]
                        for i, h in enumerate(headers):
                            nested.cell(0, i).text = h
                        for gen in generators:
                            r = nested.add_row().cells
                            r[0].text = str(gen['Generatore'])
                            r[1].text = gen['Marca e Modello']
                            r[2].text = gen['Combustibile']
                            r[3].text = str(gen['Potenza (kW)'])
                        return

    # Replace placeholders in all paragraphs and tables
    replace_placeholders_in_doc(doc, {k: v for k, v in data.items() if k != "GENERATORI"})

    # Handle conditional content
    conditional_blocks(doc, data.get("TOT_POT", 0))

    # Add the generator data table
    insert_generator_table(doc, data["GENERATORI"], "{{TAB_DATA}}")

    return doc
